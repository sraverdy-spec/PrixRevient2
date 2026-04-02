"""Generate PDF and DOCX from Markdown documentation files."""
import re
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

DOCS_DIR = "/app/docs"

# ============================================================
# DOCX GENERATION
# ============================================================

def md_to_docx(md_path, docx_path, title):
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.space_after = Pt(6)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Skip horizontal rules
        if line.strip() in ('---', '***', '___'):
            i += 1
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue
        
        if in_code_block:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(30, 30, 30)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue
        
        # Table detection
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows
            if all(set(c.strip()).issubset({'-', ':', ' '}) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # Check if next line is NOT a table row
            if i + 1 >= len(lines) or '|' not in lines[i+1] or not lines[i+1].strip().startswith('|'):
                # Flush table
                if table_rows:
                    max_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                    tbl.style = 'Light Grid Accent 1'
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_text in enumerate(row_data):
                            if ci < max_cols:
                                cell = tbl.cell(ri, ci)
                                cell.text = cell_text.strip()
                                for paragraph in cell.paragraphs:
                                    paragraph.style = doc.styles['Normal']
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)
                                        if ri == 0:
                                            run.bold = True
                    doc.add_paragraph()
                    table_rows = []
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level <= 4:
                doc.add_heading(text, level=min(level, 4))
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            i += 1
            continue
        
        # Bold/italic processing
        text = line.strip()
        p = doc.add_paragraph()
        
        # Process inline formatting
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 47, 167)
            else:
                p.add_run(part)
        
        i += 1
    
    doc.save(docx_path)
    print(f"  DOCX: {docx_path}")


# ============================================================
# PDF GENERATION
# ============================================================

class DocPDF(FPDF):
    def __init__(self, doc_title):
        super().__init__()
        self.doc_title = doc_title
        self.add_font('Liberation', '', '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf')
        self.add_font('Liberation', 'B', '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf')
        self.add_font('Liberation', 'I', '/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf')
        self.add_font('LiberationMono', '', '/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf')
    
    def header(self):
        if self.page_no() > 1:
            self.set_font('Liberation', 'I', 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, self.doc_title, new_x="RIGHT", new_y="TOP")
            self.cell(0, 8, f'Page {self.page_no()}', new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(200, 200, 200)
            self.line(10, 15, 200, 15)
            self.ln(5)
    
    def footer(self):
        self.set_y(-15)
        self.set_font('Liberation', 'I', 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, 'PrixRevient - Documentation', align='C', new_x="LMARGIN", new_y="NEXT")


def md_to_pdf(md_path, pdf_path, title):
    pdf = DocPDF(title)
    pdf.set_auto_page_break(auto=True, margin=20)
    
    # Cover page
    pdf.add_page()
    pdf.ln(60)
    pdf.set_font('Liberation', 'B', 28)
    pdf.set_text_color(0, 47, 167)
    pdf.multi_cell(0, 14, title, align='C')
    pdf.ln(10)
    pdf.set_font('Liberation', '', 14)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, 'PrixRevient - Calculateur de Prix de Revient', align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    pdf.set_font('Liberation', 'I', 11)
    pdf.cell(0, 10, 'Document généré le 2 avril 2026', align='C', new_x="LMARGIN", new_y="NEXT")
    
    # Content
    pdf.add_page()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Skip horizontal rules
        if line.strip() in ('---', '***', '___'):
            pdf.ln(3)
            i += 1
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                pdf.ln(2)
            i += 1
            continue
        
        if in_code_block:
            pdf.set_x(10)
            pdf.set_font('LiberationMono', '', 8)
            pdf.set_fill_color(245, 245, 245)
            pdf.set_text_color(30, 30, 30)
            text = line if line.strip() else ' '
            pdf.multi_cell(0, 5, text, fill=True)
            i += 1
            continue
        
        # Tables
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(set(c.strip()).issubset({'-', ':', ' '}) for c in cells):
                i += 1
                continue
            
            # Collect all table rows
            table_data = [cells]
            j = i + 1
            while j < len(lines) and '|' in lines[j] and lines[j].strip().startswith('|'):
                row_cells = [c.strip() for c in lines[j].strip().strip('|').split('|')]
                if not all(set(c.strip()).issubset({'-', ':', ' '}) for c in row_cells):
                    table_data.append(row_cells)
                j += 1
            
            # Render table
            num_cols = max(len(r) for r in table_data)
            col_width = min((190 - 2) / max(num_cols, 1), 60)
            total_w = col_width * num_cols
            start_x = 10 + (190 - total_w) / 2
            
            for ri, row in enumerate(table_data):
                pdf.set_x(start_x)
                for ci in range(num_cols):
                    cell_text = row[ci].strip() if ci < len(row) else ''
                    # Clean markdown
                    cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                    cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
                    # Truncate if too long
                    if len(cell_text) > 50:
                        cell_text = cell_text[:47] + '...'
                    
                    if ri == 0:
                        pdf.set_font('Liberation', 'B', 8)
                        pdf.set_fill_color(0, 47, 167)
                        pdf.set_text_color(255, 255, 255)
                    else:
                        pdf.set_font('Liberation', '', 8)
                        if ri % 2 == 0:
                            pdf.set_fill_color(250, 250, 250)
                        else:
                            pdf.set_fill_color(255, 255, 255)
                        pdf.set_text_color(30, 30, 30)
                    
                    pdf.cell(col_width, 7, cell_text, border=1, fill=True)
                pdf.ln()
            
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(10)
            pdf.ln(3)
            i = j
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line.split(' ')[0])
            text = line.lstrip('#').strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            
            if level == 1:
                pdf.ln(8)
                pdf.set_font('Liberation', 'B', 20)
                pdf.set_text_color(0, 47, 167)
                pdf.multi_cell(0, 10, text)
                pdf.set_draw_color(0, 47, 167)
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(5)
            elif level == 2:
                pdf.ln(6)
                pdf.set_font('Liberation', 'B', 15)
                pdf.set_text_color(0, 47, 167)
                pdf.multi_cell(0, 8, text)
                pdf.ln(2)
            elif level == 3:
                pdf.ln(4)
                pdf.set_font('Liberation', 'B', 12)
                pdf.set_text_color(50, 50, 50)
                pdf.multi_cell(0, 7, text)
                pdf.ln(1)
            else:
                pdf.ln(3)
                pdf.set_font('Liberation', 'B', 11)
                pdf.set_text_color(80, 80, 80)
                pdf.multi_cell(0, 6, text)
            
            pdf.set_text_color(0, 0, 0)
            i += 1
            continue
        
        # Empty lines
        if not line.strip():
            pdf.ln(3)
            i += 1
            continue
        
        # Normal text
        text = line.strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        pdf.set_x(10)
        pdf.set_font('Liberation', '', 10)
        pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 6, text)
        
        i += 1
    
    pdf.output(pdf_path)
    print(f"  PDF:  {pdf_path}")


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    files = [
        ("GUIDE_MISE_A_JOUR_VPS", "Guide de Mise à Jour VPS"),
        ("DOCUMENTATION_METIER", "Documentation Métier"),
        ("MODELE_DONNEES", "Modèle de Données"),
    ]
    
    for base, title in files:
        md = os.path.join(DOCS_DIR, f"{base}.md")
        print(f"\nGeneration: {title}")
        md_to_pdf(md, os.path.join(DOCS_DIR, f"{base}.pdf"), title)
        md_to_docx(md, os.path.join(DOCS_DIR, f"{base}.docx"), title)
    
    print("\nTermine ! Fichiers generes dans /app/docs/")
